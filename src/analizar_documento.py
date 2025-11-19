#!/usr/bin/env python3
"""
Script para analizar el documento Word y extraer su estructura
"""
from docx import Document
import sys

def analizar_documento(ruta_archivo):
    """Analiza un documento Word y muestra su estructura"""
    try:
        doc = Document(ruta_archivo)

        print("=" * 80)
        print("ANÁLISIS DEL DOCUMENTO WORD")
        print("=" * 80)
        print()

        # Información básica
        print(f"📄 Archivo: {ruta_archivo}")
        print(f"📝 Número de párrafos: {len(doc.paragraphs)}")
        print(f"📊 Número de tablas: {len(doc.tables)}")
        print()

        # Extraer texto de los primeros párrafos
        print("=" * 80)
        print("CONTENIDO DEL DOCUMENTO (primeros 50 párrafos)")
        print("=" * 80)
        print()

        for i, para in enumerate(doc.paragraphs[:50], 1):
            texto = para.text.strip()
            if texto:
                estilo = para.style.name if para.style else "Normal"
                print(f"[Párrafo {i}] [{estilo}]")
                print(f"  {texto[:200]}")  # Primeros 200 caracteres
                print()

        # Analizar tablas
        if doc.tables:
            print("=" * 80)
            print("TABLAS EN EL DOCUMENTO")
            print("=" * 80)
            print()

            for i, tabla in enumerate(doc.tables, 1):
                print(f"Tabla {i}:")
                print(f"  - Filas: {len(tabla.rows)}")
                print(f"  - Columnas: {len(tabla.columns) if tabla.rows else 0}")

                # Mostrar primeras 3 filas de cada tabla
                for j, fila in enumerate(tabla.rows[:3], 1):
                    celdas = [celda.text.strip() for celda in fila.cells]
                    print(f"  Fila {j}: {' | '.join(celdas[:5])}")
                print()

        # Buscar campos que parezcan variables
        print("=" * 80)
        print("POSIBLES CAMPOS A AUTOMATIZAR")
        print("=" * 80)
        print()

        campos_encontrados = set()
        for para in doc.paragraphs:
            texto = para.text
            # Buscar patrones comunes
            if any(keyword in texto.lower() for keyword in ['proyecto', 'cliente', 'fecha', 'nombre', 'dirección', 'código', 'número']):
                campos_encontrados.add(texto.strip()[:100])

        for campo in list(campos_encontrados)[:20]:
            print(f"  • {campo}")

        print()
        print("=" * 80)

    except Exception as e:
        print(f"❌ Error al analizar el documento: {e}")
        sys.exit(1)

if __name__ == "__main__":
    if len(sys.argv) > 1:
        analizar_documento(sys.argv[1])
    else:
        analizar_documento("9393653-Memoria_tecvf.docx")
