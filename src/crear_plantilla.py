#!/usr/bin/env python3
"""
Script auxiliar para crear plantillas con marcadores desde documentos existentes
Proporciona guía para convertir un documento Word en una plantilla automatizable
"""
from docx import Document
import sys
import os


def analizar_y_sugerir_marcadores(ruta_documento):
    """
    Analiza un documento y sugiere qué texto podría convertirse en marcadores

    Args:
        ruta_documento (str): Ruta al documento Word
    """
    try:
        doc = Document(ruta_documento)

        print("=" * 80)
        print("ANÁLISIS DEL DOCUMENTO PARA CREAR PLANTILLA")
        print("=" * 80)
        print()
        print(f"📄 Documento: {ruta_documento}")
        print()

        print("=" * 80)
        print("GUÍA PARA CREAR PLANTILLA")
        print("=" * 80)
        print()
        print("Para convertir este documento en una plantilla automatizable:")
        print()
        print("1. Abra el documento en Microsoft Word")
        print()
        print("2. Identifique los datos que varían en cada documento, por ejemplo:")
        print("   - Nombres de empresa")
        print("   - RUT o identificadores")
        print("   - Direcciones")
        print("   - Fechas")
        print("   - Números (personal, superficies, etc.)")
        print("   - Descripciones específicas")
        print()
        print("3. Reemplace esos datos con marcadores en formato: {{NOMBRE_MARCADOR}}")
        print()
        print("   Ejemplos:")
        print("   - 'DAVID BROWN SANTASALO' → {{EMPRESA}}")
        print("   - '76.130.276-0' → {{RUT}}")
        print("   - 'Pudahuel Poniente 1107' → {{DIRECCION}}")
        print("   - '36 trabajadores' → {{PERSONAL}} trabajadores")
        print("   - '7.741,06 m²' → {{SUPERFICIE_TOTAL}} m²")
        print()
        print("4. Guarde el documento modificado como 'plantilla_nombre.docx'")
        print()
        print("5. Use la aplicación para cargar la plantilla y completar los campos")
        print()

        print("=" * 80)
        print("CONTENIDO ANALIZADO DEL DOCUMENTO")
        print("=" * 80)
        print()

        # Buscar en tablas (suelen contener datos importantes)
        if doc.tables:
            print("📊 TABLAS ENCONTRADAS:")
            print()
            for i, tabla in enumerate(doc.tables, 1):
                print(f"Tabla {i}:")
                for j, fila in enumerate(tabla.rows[:5], 1):  # Primeras 5 filas
                    celdas = [celda.text.strip() for celda in fila.cells if celda.text.strip()]
                    if celdas:
                        print(f"  Fila {j}: {' | '.join(celdas[:3])}")
                print()

        # Buscar párrafos con datos relevantes
        print("📝 PÁRRAFOS CON POSIBLES DATOS VARIABLES:")
        print()

        palabras_clave = ['empresa', 'rut', 'dirección', 'comuna', 'región',
                         'personal', 'trabajadores', 'superficie', 'proyecto',
                         'cliente', 'fecha', 'nombre', 'código', 'número']

        parrafos_relevantes = []
        for para in doc.paragraphs:
            texto = para.text.strip()
            if texto and any(palabra in texto.lower() for palabra in palabras_clave):
                if len(texto) < 300:  # Evitar párrafos muy largos
                    parrafos_relevantes.append(texto)

        for i, texto in enumerate(parrafos_relevantes[:15], 1):
            print(f"{i}. {texto}")
            print()

        print("=" * 80)
        print("RECOMENDACIONES")
        print("=" * 80)
        print()
        print("✓ Use nombres descriptivos para los marcadores (en mayúsculas)")
        print("✓ Use guiones bajos para separar palabras: {{SUPERFICIE_TOTAL}}")
        print("✓ Sea consistente con los nombres de los marcadores")
        print("✓ Puede crear múltiples plantillas para diferentes tipos de documentos")
        print("✓ Guarde una copia del documento original antes de modificarlo")
        print()

        # Sugerir marcadores comunes
        print("=" * 80)
        print("MARCADORES SUGERIDOS PARA ESTE TIPO DE DOCUMENTO")
        print("=" * 80)
        print()

        marcadores_sugeridos = [
            ("{{EMPRESA}}", "Nombre de la empresa"),
            ("{{RUT}}", "RUT de la empresa"),
            ("{{DIRECCION}}", "Dirección principal"),
            ("{{COMUNA}}", "Comuna"),
            ("{{REGION}}", "Región"),
            ("{{PERSONAL}}", "Número de trabajadores"),
            ("{{HORARIO}}", "Horario de trabajo"),
            ("{{SUPERFICIE_TOTAL}}", "Superficie total del terreno"),
            ("{{SUPERFICIE_CONSTRUIDA}}", "Superficie construida"),
            ("{{ACTIVIDAD}}", "Descripción de la actividad"),
            ("{{FECHA}}", "Fecha del documento"),
        ]

        for marcador, descripcion in marcadores_sugeridos:
            print(f"  {marcador:30} - {descripcion}")

        print()
        print("=" * 80)

    except FileNotFoundError:
        print(f"❌ Error: No se encontró el archivo '{ruta_documento}'")
    except Exception as e:
        print(f"❌ Error al analizar documento: {e}")


def crear_plantilla_ejemplo():
    """Crea una plantilla de ejemplo con marcadores comunes"""
    print("=" * 80)
    print("CREANDO PLANTILLA DE EJEMPLO")
    print("=" * 80)
    print()

    try:
        doc = Document()

        # Título
        titulo = doc.add_heading('MEMORIA DESCRIPTIVA DE PROCESOS Y ACTIVIDADES', 0)

        # Información de empresa
        doc.add_heading('INFORMACIÓN DE LA EMPRESA', 1)

        tabla = doc.add_table(rows=5, cols=2)
        tabla.style = 'Light Grid Accent 1'

        datos_tabla = [
            ('EMPRESA', '{{EMPRESA}}'),
            ('RUT', '{{RUT}}'),
            ('DIRECCIÓN', '{{DIRECCION}}'),
            ('COMUNA', '{{COMUNA}}'),
            ('REGIÓN', '{{REGION}}'),
        ]

        for i, (campo, valor) in enumerate(datos_tabla):
            tabla.rows[i].cells[0].text = campo
            tabla.rows[i].cells[1].text = valor

        # Sección de personal
        doc.add_heading('PERSONAL Y HORARIOS', 1)
        doc.add_paragraph(
            f'El número de personal es de {{{{PERSONAL}}}} trabajadores. '
            f'El horario de trabajo es {{{{HORARIO}}}}.'
        )

        # Sección de superficies
        doc.add_heading('SUPERFICIES', 1)
        doc.add_paragraph(
            f'El terreno cuenta con una superficie total de {{{{SUPERFICIE_TOTAL}}}} m², '
            f'de los cuales {{{{SUPERFICIE_CONSTRUIDA}}}} m² corresponden a superficie construida.'
        )

        # Sección de actividad
        doc.add_heading('DESCRIPCIÓN DE ACTIVIDADES', 1)
        doc.add_paragraph('{{DESCRIPCION_ACTIVIDAD}}')

        # Guardar
        ruta_salida = 'templates/plantilla_ejemplo.docx'
        os.makedirs('templates', exist_ok=True)
        doc.save(ruta_salida)

        print(f"✓ Plantilla de ejemplo creada exitosamente:")
        print(f"  {os.path.abspath(ruta_salida)}")
        print()
        print("Esta plantilla contiene los siguientes marcadores:")
        print("  - {{EMPRESA}}")
        print("  - {{RUT}}")
        print("  - {{DIRECCION}}")
        print("  - {{COMUNA}}")
        print("  - {{REGION}}")
        print("  - {{PERSONAL}}")
        print("  - {{HORARIO}}")
        print("  - {{SUPERFICIE_TOTAL}}")
        print("  - {{SUPERFICIE_CONSTRUIDA}}")
        print("  - {{DESCRIPCION_ACTIVIDAD}}")
        print()
        print("Puede usar esta plantilla como base o analizarla para entender")
        print("cómo crear sus propias plantillas.")
        print()

    except Exception as e:
        print(f"❌ Error al crear plantilla: {e}")


def main():
    """Función principal"""
    print()
    print("=" * 80)
    print("    ASISTENTE PARA CREAR PLANTILLAS")
    print("=" * 80)
    print()
    print("Este script le ayudará a convertir sus documentos Word en plantillas")
    print("automatizables.")
    print()
    print("Opciones:")
    print("  1. Analizar un documento existente")
    print("  2. Crear plantilla de ejemplo")
    print("  3. Salir")
    print()

    while True:
        opcion = input("Seleccione una opción (1-3): ").strip()

        if opcion == '1':
            print()
            ruta = input("Ingrese la ruta del documento a analizar: ").strip()
            if ruta:
                analizar_y_sugerir_marcadores(ruta)
            break

        elif opcion == '2':
            print()
            crear_plantilla_ejemplo()
            break

        elif opcion == '3':
            print("\nHasta luego!")
            break

        else:
            print("Opción no válida. Intente nuevamente.")


if __name__ == "__main__":
    if len(sys.argv) > 1:
        # Si se pasa un archivo como argumento
        analizar_y_sugerir_marcadores(sys.argv[1])
    else:
        main()
