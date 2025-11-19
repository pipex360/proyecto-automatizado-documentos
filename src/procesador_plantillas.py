#!/usr/bin/env python3
"""
Módulo para procesar plantillas de documentos Word
Permite reemplazar marcadores en formato {{CAMPO}} con valores dinámicos
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.xmlchemy import OxmlElement
import re
import os
from datetime import datetime


class ProcesadorPlantillas:
    """Procesa plantillas de Word reemplazando marcadores con valores"""

    def __init__(self, ruta_plantilla):
        """
        Inicializa el procesador con una plantilla

        Args:
            ruta_plantilla (str): Ruta al archivo .docx de la plantilla
        """
        if not os.path.exists(ruta_plantilla):
            raise FileNotFoundError(f"No se encontró la plantilla: {ruta_plantilla}")

        self.ruta_plantilla = ruta_plantilla
        self.documento = None
        self.marcadores_encontrados = set()

    def cargar_documento(self):
        """Carga el documento Word en memoria"""
        try:
            self.documento = Document(self.ruta_plantilla)
            return True
        except Exception as e:
            print(f"Error al cargar documento: {e}")
            return False

    def extraer_marcadores(self):
        """
        Extrae todos los marcadores del documento en formato {{CAMPO}}

        Returns:
            set: Conjunto de marcadores encontrados
        """
        if not self.documento:
            self.cargar_documento()

        marcadores = set()
        patron = r'\{\{([A-Z_]+)\}\}'

        # Buscar en párrafos
        for para in self.documento.paragraphs:
            encontrados = re.findall(patron, para.text)
            marcadores.update(encontrados)

        # Buscar en tablas
        for tabla in self.documento.tables:
            for fila in tabla.rows:
                for celda in fila.cells:
                    for para in celda.paragraphs:
                        encontrados = re.findall(patron, para.text)
                        marcadores.update(encontrados)

        self.marcadores_encontrados = marcadores
        return marcadores

    def reemplazar_en_parrafo(self, parrafo, reemplazos):
        """
        Reemplaza marcadores en un párrafo manteniendo el formato

        Args:
            parrafo: Objeto párrafo de python-docx
            reemplazos (dict): Diccionario con los reemplazos {marcador: valor}
        """
        texto_completo = parrafo.text

        # Buscar marcadores en el texto
        for marcador, valor in reemplazos.items():
            patron = f'{{{{{marcador}}}}}'
            if patron in texto_completo:
                texto_completo = texto_completo.replace(patron, str(valor))

        # Si hubo cambios, actualizar el párrafo
        if texto_completo != parrafo.text:
            # Guardar el formato del primer run
            formato_original = None
            if parrafo.runs:
                run_original = parrafo.runs[0]
                formato_original = {
                    'bold': run_original.bold,
                    'italic': run_original.italic,
                    'underline': run_original.underline,
                    'font_name': run_original.font.name,
                    'font_size': run_original.font.size,
                }

            # Limpiar el párrafo
            for run in parrafo.runs:
                run.text = ''

            # Agregar el nuevo texto
            nuevo_run = parrafo.add_run(texto_completo)

            # Aplicar formato original si existe
            if formato_original:
                if formato_original['bold'] is not None:
                    nuevo_run.bold = formato_original['bold']
                if formato_original['italic'] is not None:
                    nuevo_run.italic = formato_original['italic']
                if formato_original['underline'] is not None:
                    nuevo_run.underline = formato_original['underline']
                if formato_original['font_name']:
                    nuevo_run.font.name = formato_original['font_name']
                if formato_original['font_size']:
                    nuevo_run.font.size = formato_original['font_size']

    def reemplazar_marcadores(self, datos):
        """
        Reemplaza todos los marcadores del documento con los datos proporcionados

        Args:
            datos (dict): Diccionario con los valores {MARCADOR: valor}

        Returns:
            bool: True si fue exitoso, False en caso contrario
        """
        if not self.documento:
            self.cargar_documento()

        try:
            # Reemplazar en párrafos
            for para in self.documento.paragraphs:
                self.reemplazar_en_parrafo(para, datos)

            # Reemplazar en tablas
            for tabla in self.documento.tables:
                for fila in tabla.rows:
                    for celda in fila.cells:
                        for para in celda.paragraphs:
                            self.reemplazar_en_parrafo(para, datos)

            return True
        except Exception as e:
            print(f"Error al reemplazar marcadores: {e}")
            return False

    def guardar_documento(self, ruta_salida):
        """
        Guarda el documento procesado

        Args:
            ruta_salida (str): Ruta donde guardar el documento

        Returns:
            bool: True si fue exitoso, False en caso contrario
        """
        try:
            # Crear directorio si no existe
            directorio = os.path.dirname(ruta_salida)
            if directorio and not os.path.exists(directorio):
                os.makedirs(directorio)

            self.documento.save(ruta_salida)
            return True
        except Exception as e:
            print(f"Error al guardar documento: {e}")
            return False

    def procesar(self, datos, ruta_salida):
        """
        Proceso completo: cargar, reemplazar y guardar

        Args:
            datos (dict): Diccionario con los valores {MARCADOR: valor}
            ruta_salida (str): Ruta donde guardar el documento procesado

        Returns:
            bool: True si fue exitoso, False en caso contrario
        """
        if not self.cargar_documento():
            return False

        if not self.reemplazar_marcadores(datos):
            return False

        if not self.guardar_documento(ruta_salida):
            return False

        return True


def generar_nombre_archivo(prefijo="documento", extension="docx"):
    """
    Genera un nombre de archivo único con timestamp

    Args:
        prefijo (str): Prefijo del nombre
        extension (str): Extensión del archivo

    Returns:
        str: Nombre de archivo generado
    """
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    return f"{prefijo}_{timestamp}.{extension}"


# Ejemplo de uso
if __name__ == "__main__":
    # Datos de ejemplo
    datos_ejemplo = {
        "EMPRESA": "DAVID BROWN SANTASALO SOUTH AMERICA S.A",
        "RUT": "76.130.276-0",
        "DIRECCION": "Pudahuel Poniente 1107",
        "COMUNA": "Pudahuel",
        "REGION": "Región Metropolitana",
        "PERSONAL": "36",
        "HORARIO": "Lunes a Viernes, diurno",
        "SUPERFICIE_TOTAL": "7.741,06",
        "SUPERFICIE_CONSTRUIDA": "2.356,53",
    }

    # Procesar plantilla
    procesador = ProcesadorPlantillas("templates/plantilla.docx")
    procesador.cargar_documento()

    # Mostrar marcadores encontrados
    marcadores = procesador.extraer_marcadores()
    print("Marcadores encontrados:")
    for marcador in sorted(marcadores):
        print(f"  - {{{{ {marcador} }}}}")

    # Reemplazar y guardar
    nombre_salida = generar_nombre_archivo("memoria_tecnica")
    ruta_salida = f"output/{nombre_salida}"

    if procesador.procesar(datos_ejemplo, ruta_salida):
        print(f"\n✓ Documento generado exitosamente: {ruta_salida}")
    else:
        print("\n✗ Error al procesar el documento")
